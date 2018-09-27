'''
Developed By Felipe Garcia Diaz
'''
import wikipedia
import time
from docx import Document

#Main Function
def run():
    #Init Document
    doc = Document()
    doc.add_heading("Study Sheet",0)
    #Basic Questions
    studyFile = input("What file would you like to use? >>")
    outputFile=input("Where would you like to save the file >>")
    num = input("Sentences to parse? >>")
    #I.E. Debug mode, to check for speed and errors
    verboseMode = input("Would you like to use verbose mode? [Y/N]>>").lower()
    studyMat = open(studyFile).readlines()
    #The loop, the most beautiful part, will reep through each article and beautifully create a Word Document
    for i in range(len(studyMat)):
        try:
            if verboseMode == "y" or None:
                print(studyMat[i] + "\n" + wikipedia.summary(studyMat[i], sentences=num))
            else:    
                pass
            doc.add_heading(studyMat[i], level=1)
            desc = doc.add_paragraph(wikipedia.summary(studyMat[i], sentences=num))
            #I added this to keep wikipedia from booting you
            time.sleep(5)
        except:
            print("There was an error reading " + studyMat[i] + " Ignoring for now")
            pass
    #Saves the document
    doc.add_page_break()
    doc.save(outputFile)
#All the cool python developers put this at the end of there code. So i did too. Plus it actually has some useful implementations
if __name__ == "__main__":
    run()
#TODO: Add a gui to this program, to make it more easy for everyday users.
